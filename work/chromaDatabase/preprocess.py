import os, sys
sys.path.append("..")
curr_dir = os.path.dirname(os.path.abspath(__file__))
DOCS_PATH = os.path.join(curr_dir, "documents")

import logging, preheader # import for custom logger
logger = logging.getLogger(__name__) # Setup logging

from docx import Document
from doc2docx import convert
from pypdf import PdfReader
import pandas as pd
from openpyxl import load_workbook
import re, shutil, platform, subprocess

# -----------------------------------------------------------------------------------------------------------

def delete_doc_sections(docx_path, titles_to_delete=["Contents, References, Annex"], remove_tabes=True):
    '''
    Description:
        Remove "Contents" Section: Deletes the "Contents" heading and all following paragraphs until the next heading.
        Remove "References" Section: Uses remove_content_after_heading to delete the "References" section.
        Remove Everything After "Annex": Finds the first heading containing "Annex" and deletes it and all subsequent paragraphs.
        Remove All Tables: Iterates through all tables and removes them from the document.
    '''
    doc = Document(docx_path)

    def remove_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._element = p._p = None

    def remove_content_after_heading(doc, heading_text, delete_to_next_heading=True):
        delete = False
        for paragraph in doc.paragraphs:
            if heading_text in paragraph.text and paragraph.style.name.startswith("Heading"):
                delete = True
                remove_paragraph(paragraph)
                continue
            if delete:
                if delete_to_next_heading and paragraph.style.name.startswith("Heading"):
                    break
                remove_paragraph(paragraph)

    if "Contents" in titles_to_delete:
        titles_to_delete.remove("Contents")
        delete_contents = False
        for paragraph in doc.paragraphs:
            if "Contents" in paragraph.text:
                delete_contents = True
            if delete_contents:
                if paragraph.style.name.startswith("Heading") and "Contents" not in paragraph.text:
                    break 
                remove_paragraph(paragraph)

    if "Annex" in titles_to_delete:
        titles_to_delete.remove("Annex")
        found_annex = False
        paragraphs_to_delete = []
        for i, paragraph in enumerate(doc.paragraphs):
            if "Annex" in paragraph.text and paragraph.style.name.startswith("Heading"):
                found_annex = True
            if found_annex:
                paragraphs_to_delete.append(paragraph)

        for paragraph in paragraphs_to_delete:
            remove_paragraph(paragraph)

    if remove_tabes:
        for table in doc.tables:
            tbl = table._element
            tbl.getparent().remove(tbl)

    for title in titles_to_delete: # remove other titles as specified
        remove_content_after_heading(doc, title)
    return doc

# difficult to identify block of text under each title in pdfs, so just remove contents using regrex formating
def delete_contents(raw_text):
    lines = raw_text.split('\n')
    filtered_lines = []
    for line in lines:
        # Detect the start of the Contents section
        if re.match(r'^\s*Contents\s*$', line.strip()):
            continue  # Skip the "Contents" line itself
        # 
        # Skip lines that match the Contents pattern (title + dots + page number) like "Title ......... page_number" (with lots of dots or spaces before the number)
        if re.match(r'^.+?(\.{3,}| +)\d+\s*$', line.strip()):
            continue
        filtered_lines.append(line) # Keep all other lines
    return '\n'.join(filtered_lines)

# -----------------------------------------------------------------------------------------------------------

def convert_docs_to_docx(folder_path, clear_space=False):
    """Convert .doc files in a folder to .docx format if any."""
    has_doc = any(f.endswith('.doc') for f in os.listdir(folder_path))
    if not has_doc:
        return
    if platform.system().lower() == "linux":
        # Convert all .doc files in the folder to .docx using LibreOffice
        doc_files = [os.path.join(folder_path, f) for f in os.listdir(folder_path) if f.endswith('.doc') and not f.endswith('.docx')]
        try:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", folder_path] + doc_files,
                check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
            )
        except Exception as e:
            logger.warning(f"Failed to convert DOC files using LibreOffice: {e}")
    else:
        # Use doc2docx's convert function for Windows or other OS
        try:
            convert(folder_path)
        except Exception as e:
            logger.warning(f"Failed to convert DOC files using doc2docx: {e}")

    if not clear_space:
        return
    # clear any other temp files/folders other than .docx files
    for f in os.listdir(folder_path):
        file_path = os.path.join(folder_path, f)
        if f.endswith('.docx'):
            continue
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.remove(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            logger.warning(f"Failed to delete {file_path}: {e}")

def read_docx(file_path):
    """Read and extract text from a DOCX file."""
    text_extract = None; metadata = None
    try:
        doc = delete_doc_sections(file_path, titles_to_delete=["Contents"], remove_tabes=False) # remove only content page for now
        text_extract = '\n'.join(para.text for para in doc.paragraphs)
        
    except Exception as e:
        logger.error(f"Failed to read DOCX file at {file_path}: {e}. Removing corrupted file.")
        #os.remove(file_path) # comment to ignore
    try:
        prop = doc.core_properties # remove private or unnecessary attrb startin with '_'
        metadata = {d: getattr(prop, d) for d in dir(prop) if not d.startswith('_')}
    except Exception as e:
        logger.error(f"Failed to read DOCX file metadata! {file_path}: {e}")
        
    return text_extract, metadata

def read_pdf(file_path):
    """Read and extract text from a PDF file."""
    text_extract = None; metadata = None
    try:
        pdf = PdfReader(file_path)
        raw_text = '\n'.join(page.extract_text() for page in pdf.pages)
        text_extract = delete_contents(raw_text)
        #logger.info(text_extract)
    except Exception as e:
        logger.error(f"Failed to read PDF file at {file_path}: {e}. Removing corrupted file.")
        #os.remove(file_path) # comment to ignore
    try:
        prop = pdf.metadata #strip the leading '/' from all string keys in the metadata.
        metadata = {key.lstrip('/') if isinstance(key, str) else key: value for key, value in prop.items()}
    except Exception as e:
        logger.error(f"Failed to read DOCX file metadata! {file_path}: {e}")
    return text_extract, metadata

def read_xcel(file_path):
    """Read and extract list of texts from a XLSX or CSV file."""
    text_extract = None; metadata = None
    try:
        # Determine file type
        if file_path.endswith('.xlsx'):
            df = pd.read_excel(file_path)
        elif file_path.endswith('.csv'):
            df = pd.read_csv(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path}")
        
        # Capture original DataFrame properties
        original_num_columns = len(df.columns)
        original_missing = df.isna().sum().to_dict()
        all_na_cols = [col for col in df.columns if df[col].isna().all()]
        num_dropped_columns = len(all_na_cols)

        # Drop columns with all missing values
        df = df.dropna(axis=1, how='all')
        # Drop columns with only numeric values (all values are numbers or NaN)
        numeric_only_cols = [
            col for col in df.columns
            if pd.api.types.is_numeric_dtype(df[col]) and df[col].notna().all()
        ]
        df = df.drop(columns=numeric_only_cols)
        # Process each row to generate enriched strings
        text_extract = []
        for _, row in df.iterrows():
            attributes = []
            for col, val in row.items():
                if pd.notna(val) and val != '':  # Only include valid values
                    attributes.append(f"{col}: {val}")
            if attributes:
                text_extract.append('\n'.join(attributes))

        # Add DataFrame properties (always present)
        metadata = {
            'original_num_columns': original_num_columns,
            'num_dropped_columns': num_dropped_columns,
            'dropped_columns': all_na_cols,
            'num_rows': len(df),
            'num_columns': len(df.columns),
            'columns': list(df.columns),
            'dtypes': df.dtypes.astype(str).to_dict(),
            'memory_usage': int(df.memory_usage(deep=True).sum()),
            'missing_values_before_drop': original_missing,
        }
    except Exception as e:
        logger.error(f"Failed to read XLSX/CSV file at {file_path}: {e}. Removing corrupted file.")
        #os.remove(file_path) # comment to ignore
    
    try:
        if file_path.endswith('.xlsx'):     
            wb = load_workbook(file_path, read_only=True, keep_vba=False)
            props = wb.properties
            # Extract relevant Excel properties with 'xlsx_' prefix
            file_meta = {f'xlsx_{k}': (v.isoformat() if hasattr(v, 'isoformat') else v)
                        for k, v in vars(props).items()
                        if not k.startswith('_') and v}
            wb.close()
            metadata.update(file_meta)
        else:
            file_stats = os.stat(file_path)
            file_meta = {
            'file_name': os.path.basename(file_path),
            'file_size': file_stats.st_size,
            'modified_time': file_stats.st_mtime,
            }
            metadata.update(file_meta)
    except Exception as wb_error:
        logger.error(f"Failed to read XLSX/CSV file metadata! {file_path}: {e}")

    # 3. text_extract = list of enriched string of each row, return the metadata of the file as well
    return text_extract, metadata

# -----------------------------------------------------------------------------------------------------------

# Process the documents inside the directory given
def process_documents(doc_dir=DOCS_PATH):
    doc_ds = []
    for root, _, files in os.walk(doc_dir):
        for filename in files:
            file_path = os.path.join(root, filename)
            content = None; meta = None
            
            ext = os.path.splitext(filename)[1].lower() # process each file based on extension
            if ext == '.docx':
                content, meta = read_docx(file_path)
            elif ext == '.pdf':
                content, meta = read_pdf(file_path)
            if ext in ('.xlsx', '.csv'):
                content, meta = read_xcel(file_path)
            
            if meta: # conv all non primitive data types inside meta to string
                meta = {k: (str(v) if not isinstance(v, (str, int, float, bool)) else v) for k, v in meta.items()}
            if content: # create data_dict
                data_dict = {"text": content, "source": filename, "metadata": meta}
                doc_ds.append(data_dict)
    return doc_ds

# helper funtion for rag methodoligies
def find_terms_and_abbreviations_in_sentence(terms_dict, abbreviations_dict, sentence):
    """Finds and filters terms and abbreviations in the given sentence.
       Abbreviations are matched case-sensitively, terms case-insensitively, and longer terms are prioritized."""
    
    def preprocess(text, lowercase=True):
        """Converts text to lowercase and removes punctuation."""
        if lowercase:
            text = text.lower()
        punctuations = '''!()-[]{};:'"\,<>./?@#$%^&*_~'''
        for char in punctuations:
            text = text.replace(char, '')
        return text

    def find_and_filter_terms(terms_dict, sentence):
        """Finds terms in the given sentence, case-insensitively, and filters out shorter overlapping terms."""
        lowercase_sentence = preprocess(sentence, lowercase=True)
        
        matched_terms = {term: terms_dict[term] for term in terms_dict if preprocess(term) in lowercase_sentence}
        
        final_terms = {}
        for term in matched_terms:
            if not any(term in other and term != other for other in matched_terms):
                final_terms[term] = matched_terms[term]
                
        return final_terms

    def find_and_filter_abbreviations(abbreviations_dict, sentence):
        """Finds abbreviations in the given sentence, case-sensitively, and filters out shorter overlapping abbreviations."""
        processed_sentence = preprocess(sentence, lowercase=False)  
        words = processed_sentence.split() 
        
        matched_abbreviations = {word: abbreviations_dict[word] for word in words if word in abbreviations_dict}
        
        final_abbreviations = {}
        sorted_abbrs = sorted(matched_abbreviations, key=len, reverse=True)
        for abbr in sorted_abbrs:
            if not any(abbr in other and abbr != other for other in sorted_abbrs):
                final_abbreviations[abbr] = matched_abbreviations[abbr]
        
        return final_abbreviations
    
    matched_terms = find_and_filter_terms(terms_dict, sentence)
    matched_abbreviations = find_and_filter_abbreviations(abbreviations_dict, sentence)

    formatted_terms = [f"{term}: {definition}" for term, definition in matched_terms.items()]
    formatted_abbreviations = [f"{abbr}: {definition}" for abbr, definition in matched_abbreviations.items()]

    return formatted_terms, formatted_abbreviations
